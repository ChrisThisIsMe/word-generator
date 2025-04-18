
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import os

def generate_docx():
    md_path = filedialog.askopenfilename(filetypes=[("Markdown files", "*.md"), ("Text files", "*.txt")])
    if not md_path:
        return

    try:
        marker_keyword = entry_marker_keyword.get().strip().lower()
        blank_after_marker = int(entry_blank_after_marker.get())
        marker_found = False

        h1_size = float(entry_h1_size.get())
        h2_size = float(entry_h2_size.get())
        h3_size = float(entry_h3_size.get())
        normal_size = float(entry_normal_size.get())
        h1_before = int(entry_h1_before.get())
        h1_after = int(entry_h1_after.get())
        h2_before = int(entry_h2_before.get())
        h2_after = int(entry_h2_after.get())
        h3_before = int(entry_h3_before.get())
        h3_after = int(entry_h3_after.get())
        normal_before = int(entry_normal_before.get())
        normal_after = int(entry_normal_after.get())
        h1_spacing = float(entry_h1_spacing.get())
        h2_spacing = float(entry_h2_spacing.get())
        h3_spacing = float(entry_h3_spacing.get())
        normal_spacing = float(entry_normal_spacing.get())
        blank_pages = int(entry_blank_pages.get())

        page_width = float(entry_page_width.get())
        page_height = float(entry_page_height.get())
        margin_top = float(entry_margin_top.get())
        margin_bottom = float(entry_margin_bottom.get())
        margin_left = float(entry_margin_left.get())
        margin_right = float(entry_margin_right.get())
    except ValueError:
        messagebox.showerror("Erreur", "Veuillez entrer des valeurs valides.")
        return

    h1_font = entry_h1_font.get()
    h2_font = entry_h2_font.get()
    h3_font = entry_h3_font.get()
    normal_font = entry_normal_font.get()

    add_page_break_h1 = var_page_break_h1.get()
    add_page_break_h2 = var_page_break_h2.get()
    uppercase_h1 = var_uppercase_h1.get()
    uppercase_h2 = var_uppercase_h2.get()
    merge_lines = var_merge_lines.get()

    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(page_width)
    section.page_height = Mm(page_height)
    section.top_margin = Cm(margin_top)
    section.bottom_margin = Cm(margin_bottom)
    section.left_margin = Cm(margin_left)
    section.right_margin = Cm(margin_right)

    for _ in range(blank_pages):
        doc.add_paragraph("")
        doc.add_page_break()

    align_map = {
        "Gauche": WD_ALIGN_PARAGRAPH.LEFT,
        "Centré": WD_ALIGN_PARAGRAPH.CENTER,
        "Droite": WD_ALIGN_PARAGRAPH.RIGHT,
        "Justifié": WD_ALIGN_PARAGRAPH.JUSTIFY
    }

    def customize_style(style, font_name, font_size, spacing_before, spacing_after, line_spacing):
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.paragraph_format.space_before = Pt(spacing_before)
        style.paragraph_format.space_after = Pt(spacing_after)
        style.paragraph_format.line_spacing = line_spacing

    customize_style(doc.styles['Heading 1'], h1_font, h1_size, h1_before, h1_after, h1_spacing)
    customize_style(doc.styles['Heading 2'], h2_font, h2_size, h2_before, h2_after, h2_spacing)
    customize_style(doc.styles['Heading 3'], h3_font, h3_size, h3_before, h3_after, h3_spacing)
    customize_style(doc.styles['Normal'], normal_font, normal_size, normal_before, normal_after, normal_spacing)

    align_h1 = align_map[alignment_h1.get()]
    align_h2 = align_map[alignment_h2.get()]
    align_h3 = align_map[alignment_h3.get()]
    align_normal = align_map[alignment_normal.get()]

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        buffer = []
        for line in lines:
            stripped = line.strip()
            if stripped.lower() == marker_keyword:
                marker_found = True
                p = doc.add_paragraph(stripped, style="Heading 1")
                p.alignment = align_h1
                continue
            elif stripped.startswith("# "):
                if add_page_break_h1:
                    doc.add_page_break()
                text = stripped[2:].upper() if uppercase_h1 else stripped[2:]
                p = doc.add_paragraph(text, style="Heading 1")
                p.alignment = align_h1
            elif stripped.startswith("## "):
                if add_page_break_h2:
                    doc.add_page_break()
                text = stripped[3:].upper() if uppercase_h2 else stripped[3:]
                p = doc.add_paragraph(text, style="Heading 2")
                p.alignment = align_h2
            elif stripped.startswith("### "):
                text = stripped[4:]
                p = doc.add_paragraph(text, style="Heading 3")
                p.alignment = align_h3
            elif stripped == "":
                if buffer:
                    paragraph = " ".join(buffer).strip()
                    if paragraph:
                        p = doc.add_paragraph(paragraph, style="Normal")
                        p.alignment = align_normal
                    buffer = []
            else:
                buffer.append(stripped)

    try:
        blank_end = int(entry_blank_end.get())
    except ValueError:
        blank_end = 0
    for _ in range(blank_end):
        doc.add_paragraph("")
        doc.add_page_break()

    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word documents", "*.docx")],
        initialfile=os.path.splitext(os.path.basename(md_path))[0] + "_styled.docx",
        title="Enregistrer le document Word"
    )

    if save_path:
        doc.save(save_path)
        messagebox.showinfo("Succès", f"Document généré :\n{save_path}")

root = tk.Tk()
root.title("Generator Word Style")
root.geometry("800x700")

window_width = 800
window_height = 700
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
x_centre = int((screen_width - window_width) / 2)
y_centre = int((screen_height - window_height) / 2)
root.geometry(f"{window_width}x{window_height}+{x_centre}+{y_centre}")

root.lift()
root.attributes("-topmost", True)
root.after(100, lambda: root.attributes("-topmost", False))


frame_styles = tk.LabelFrame(root, text="Styles & Tailles")
frame_styles.grid(row=0, column=0, columnspan=4, padx=10, pady=5, sticky='we')

entry_h1_font = tk.Entry(frame_styles); entry_h1_font.insert(0, "Garamond")
entry_h2_font = tk.Entry(frame_styles); entry_h2_font.insert(0, "Garamond")
entry_h3_font = tk.Entry(frame_styles); entry_h3_font.insert(0, "Garamond")
entry_normal_font = tk.Entry(frame_styles); entry_normal_font.insert(0, "Garamond")

entry_h1_size = tk.Entry(frame_styles, width=5); entry_h1_size.insert(0, "24")
entry_h2_size = tk.Entry(frame_styles, width=5); entry_h2_size.insert(0, "16")
entry_h3_size = tk.Entry(frame_styles, width=5); entry_h3_size.insert(0, "13")
entry_normal_size = tk.Entry(frame_styles, width=5); entry_normal_size.insert(0, "11")

labels = ["Heading 1 (police)", "Heading 2 (police)", "Heading 3 (police)", "Normal (police)"]
entries = [(entry_h1_font, entry_h1_size), (entry_h2_font, entry_h2_size), (entry_h3_font, entry_h3_size), (entry_normal_font, entry_normal_size)]

for i, (label, (font_entry, size_entry)) in enumerate(zip(labels, entries)):
    tk.Label(frame_styles, text=label).grid(row=i, column=0)
    font_entry.grid(row=i, column=1)
    tk.Label(frame_styles, text="Taille").grid(row=i, column=2)
    size_entry.grid(row=i, column=3)

frame_spacing = tk.LabelFrame(root, text="Espacements & Interligne")
frame_spacing.grid(row=1, column=0, columnspan=4, padx=10, pady=5, sticky='we')

spacing_data = [
    ("H1", "0", "12", "2.0"),
    ("H2", "0", "9", "1.15"),
    ("H3", "0", "9", "1.15"),
    ("Normal", "0", "6", "1.15")
]
entries_spacing = []

for i, (label, before, after, spacing) in enumerate(spacing_data):
    tk.Label(frame_spacing, text=f"{label} avant / après").grid(row=i, column=0)
    e_before = tk.Entry(frame_spacing, width=5); e_before.insert(0, before)
    e_after = tk.Entry(frame_spacing, width=5); e_after.insert(0, after)
    e_spacing = tk.Entry(frame_spacing, width=5); e_spacing.insert(0, spacing)
    e_before.grid(row=i, column=1); e_after.grid(row=i, column=2)
    tk.Label(frame_spacing, text=f"Interligne {label}").grid(row=i, column=3)
    e_spacing.grid(row=i, column=4)
    entries_spacing.append((e_before, e_after, e_spacing))

(entry_h1_before, entry_h1_after, entry_h1_spacing), (entry_h2_before, entry_h2_after, entry_h2_spacing), (entry_h3_before, entry_h3_after, entry_h3_spacing), (entry_normal_before, entry_normal_after, entry_normal_spacing) = entries_spacing

frame_options = tk.LabelFrame(root, text="Alignement, Majuscules, Pages")
frame_options.grid(row=2, column=0, columnspan=4, padx=10, pady=5, sticky='we')

alignment_h1 = tk.StringVar(value="Centré")
alignment_h2 = tk.StringVar(value="Gauche")
alignment_h3 = tk.StringVar(value="Gauche")
alignment_normal = tk.StringVar(value="Justifié")

alignments = [("Alignement H1", alignment_h1), ("Alignement H2", alignment_h2), ("Alignement H3", alignment_h3), ("Alignement Normal", alignment_normal)]
for i, (label, var) in enumerate(alignments):
    tk.Label(frame_options, text=label).grid(row=i, column=0, sticky='e')
    tk.OptionMenu(frame_options, var, "Gauche", "Centré", "Droite", "Justifié").grid(row=i, column=1)

var_uppercase_h1 = tk.BooleanVar(value=True)
var_uppercase_h2 = tk.BooleanVar()
var_page_break_h1 = tk.BooleanVar(value=True)
var_page_break_h2 = tk.BooleanVar(value=True)
var_merge_lines = tk.BooleanVar(value=True)

tk.Checkbutton(frame_options, text="H1 en MAJUSCULES", variable=var_uppercase_h1).grid(row=0, column=2, sticky='w')
tk.Checkbutton(frame_options, text="H2 en MAJUSCULES", variable=var_uppercase_h2).grid(row=1, column=2, sticky='w')
tk.Checkbutton(frame_options, text="Saut de page avant H1", variable=var_page_break_h1).grid(row=0, column=3, sticky='w')
tk.Checkbutton(frame_options, text="Saut de page avant H2", variable=var_page_break_h2).grid(row=1, column=3, sticky='w')
tk.Checkbutton(frame_options, text="Fusionner les lignes Markdown", variable=var_merge_lines).grid(row=2, column=2, columnspan=2, sticky='w')

entry_blank_pages = tk.Entry(frame_options, width=5); entry_blank_pages.insert(0, "2")
tk.Label(frame_options, text="Pages blanches au début").grid(row=4, column=0, sticky='e')
entry_blank_pages.grid(row=4, column=1)

entry_marker_keyword = tk.Entry(frame_options, width=20); entry_marker_keyword.insert(0, "Page Légale")
tk.Label(frame_options, text="Repère insertion pages").grid(row=5, column=0, sticky='e')
entry_marker_keyword.grid(row=5, column=1)

entry_blank_after_marker = tk.Entry(frame_options, width=5); entry_blank_after_marker.insert(0, "1")
tk.Label(frame_options, text="Pages après repère").grid(row=5, column=2, sticky='e')
entry_blank_after_marker.grid(row=5, column=3)

entry_blank_end = tk.Entry(frame_options, width=5); entry_blank_end.insert(0, "1")
tk.Label(frame_options, text="Pages blanches à la fin").grid(row=6, column=0, sticky='e')
entry_blank_end.grid(row=6, column=1)

frame_layout = tk.LabelFrame(root, text="Mise en page")
frame_layout.grid(row=3, column=0, columnspan=4, padx=10, pady=5, sticky='we')

entry_page_width = tk.Entry(frame_layout, width=7); entry_page_width.insert(0, "120")
entry_page_height = tk.Entry(frame_layout, width=7); entry_page_height.insert(0, "190")
entry_margin_top = tk.Entry(frame_layout, width=7); entry_margin_top.insert(0, "1.8")
entry_margin_bottom = tk.Entry(frame_layout, width=7); entry_margin_bottom.insert(0, "1.8")
entry_margin_left = tk.Entry(frame_layout, width=7); entry_margin_left.insert(0, "2.0")
entry_margin_right = tk.Entry(frame_layout, width=7); entry_margin_right.insert(0, "2.0")

tk.Label(frame_layout, text="Largeur (mm)").grid(row=0, column=0)
entry_page_width.grid(row=0, column=1)
tk.Label(frame_layout, text="Hauteur (mm)").grid(row=0, column=2)
entry_page_height.grid(row=0, column=3)
tk.Label(frame_layout, text="Marge haut / bas (cm)").grid(row=1, column=0)
entry_margin_top.grid(row=1, column=1)
entry_margin_bottom.grid(row=1, column=2)
tk.Label(frame_layout, text="Marge gauche / droite (cm)").grid(row=2, column=0)
entry_margin_left.grid(row=2, column=1)
entry_margin_right.grid(row=2, column=2)

tk.Button(root, text="Choisir un fichier .md et générer le Word", command=generate_docx).grid(row=4, column=0, columnspan=4, pady=10)

root.mainloop()
